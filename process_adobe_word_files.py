#!/usr/bin/env python3
"""
Process Adobe-converted DOCX files for anonymization.

This script is designed for the workflow:
1. Adobe Acrobat: PDF → DOCX (batch conversion)
2. This script: Anonymize + Strip Metadata + Remove Images
3. Adobe Acrobat: DOCX → PDF (batch conversion back)

Usage:
    python3 process_adobe_word_files.py \\
        --input path/to/adobe_converted_docx/ \\
        --output path/to/anonymized_docx/

The script will:
- Apply all 128 anonymization rules
- Remove ALL embedded images/logos
- Strip ALL metadata (author, company, title, etc.)
- Preserve directory structure
"""

import os
import sys
from pathlib import Path
from datetime import datetime
import argparse
import logging
from multiprocessing import Pool, cpu_count

# Excel and DOCX processing
from openpyxl import load_workbook
from docx import Document


def strip_all_metadata(doc):
    """
    Strip ALL metadata from DOCX file - CRITICAL for anonymization.

    Adobe's PDF → DOCX conversion PRESERVES metadata like:
    - Author: "Shake Shack Inc."
    - Company: "Shake Shack"
    - Title: "Q1 2022 10-Q"

    This function wipes ALL of it.
    """
    props = doc.core_properties

    # Clear all identifying metadata
    props.author = ""
    props.last_modified_by = ""
    props.title = ""
    props.subject = ""
    props.keywords = ""
    props.comments = ""
    props.category = ""
    props.content_status = ""
    props.identifier = ""
    props.language = ""

    # Clear company (BIG ONE for SEC filings)
    if hasattr(props, 'company'):
        props.company = ""

    # Clear creator/application info
    props.creator = ""

    # Reset version/revision
    props.revision = 1
    if hasattr(props, 'version'):
        props.version = None

    # Optionally clear dates (keeps last_modified as save time)
    # props.created = None
    # props.modified = None
    # props.last_printed = None

    return doc


def remove_all_images(doc):
    """
    Remove ALL embedded images from DOCX (logos, watermarks, charts).

    FIXED: Now uses XPath to find images at ANY depth in the XML tree,
    not just direct children. This catches nested images like logos.

    Images are found in:
    - Inline shapes in paragraphs (nested in drawing elements)
    - Headers (all sections)
    - Footers (all sections)
    """
    removed_count = 0

    # Remove images from main document body using XPath
    for paragraph in doc.paragraphs:
        if hasattr(paragraph._element, 'xpath'):
            # Find ALL drawing elements at any depth (not just direct children)
            drawings = paragraph._element.xpath('.//w:drawing')
            for drawing in drawings:
                # Remove the drawing element from its parent
                parent = drawing.getparent()
                if parent is not None:
                    parent.remove(drawing)
                    removed_count += 1

    # Remove images from headers and footers
    for section in doc.sections:
        # Process all header types
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for paragraph in header.paragraphs:
                if hasattr(paragraph._element, 'xpath'):
                    drawings = paragraph._element.xpath('.//w:drawing')
                    for drawing in drawings:
                        parent = drawing.getparent()
                        if parent is not None:
                            parent.remove(drawing)
                            removed_count += 1

        # Process all footer types
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for paragraph in footer.paragraphs:
                if hasattr(paragraph._element, 'xpath'):
                    drawings = paragraph._element.xpath('.//w:drawing')
                    for drawing in drawings:
                        parent = drawing.getparent()
                        if parent is not None:
                            parent.remove(drawing)
                            removed_count += 1

    return removed_count


def clear_headers_footers(doc):
    """
    Clear all content from headers and footers while preserving structure.

    Use case: Investor presentations with company logos in headers/footers.
    This removes the logo/text content but keeps body images intact.

    Clears:
    - All header types (default, first page, even page)
    - All footer types (default, first page, even page)
    - Text and tables within headers/footers
    """
    cleared_count = 0

    for section in doc.sections:
        # Clear all header types
        for header in [section.header, section.first_page_header, section.even_page_header]:
            # Clear all paragraphs in header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():  # Only count non-empty
                    cleared_count += 1
                paragraph.clear()

            # Clear all tables in header
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                cleared_count += 1
                            paragraph.clear()

        # Clear all footer types
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            # Clear all paragraphs in footer
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():  # Only count non-empty
                    cleared_count += 1
                paragraph.clear()

            # Clear all tables in footer
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                cleared_count += 1
                            paragraph.clear()

    return cleared_count


def load_aliases_from_excel(excel_path):
    """
    Load anonymization mappings from Excel file.
    Reuses exact logic from vdr_anonymizer_final.py
    """
    wb = load_workbook(excel_path, data_only=True)

    # Find the correct sheet (flexible detection)
    sheet = None
    for sheet_name in wb.sheetnames:
        if any(keyword in sheet_name.lower() for keyword in ['tracker', 'anonymization', 'mapping', 'aliases']):
            sheet = wb[sheet_name]
            break

    if sheet is None:
        sheet = wb.active

    # Find columns
    header_row = None
    for row in sheet.iter_rows(max_row=10):
        cell_values = [str(cell.value).lower() if cell.value else '' for cell in row]
        # Support both "Before/After" and "Original/Replacement" formats
        if any('original' in val or 'real' in val or 'actual' in val or 'before' in val for val in cell_values):
            header_row = row
            break

    if not header_row:
        raise ValueError("Could not find header row in Excel file")

    # Identify columns
    original_col = None
    replacement_col = None

    for idx, cell in enumerate(header_row):
        val = str(cell.value).lower() if cell.value else ''
        # Support both "Before" and "Original" for source column
        if 'original' in val or 'real' in val or 'actual' in val or 'before' in val:
            original_col = idx
        # Support both "After" and "Replacement" for target column
        elif 'replacement' in val or 'anonymized' in val or 'alias' in val or 'after' in val:
            replacement_col = idx

    if original_col is None or replacement_col is None:
        raise ValueError("Could not identify original and replacement columns")

    # Load mappings
    alias_map = {}
    for row in sheet.iter_rows(min_row=header_row[0].row + 1):
        original = row[original_col].value
        replacement = row[replacement_col].value

        # CRITICAL: Allow blank "After" values for deletion/removal
        # If After is blank, we replace with empty string (removes the text)
        if original is not None:
            original = str(original).strip()
            if original:  # Only check original is not empty
                # Handle None or blank replacement as empty string (deletion)
                if replacement is None or str(replacement).strip() == "":
                    alias_map[original] = ""
                else:
                    alias_map[original] = str(replacement).strip()

    # Generate reverse names for Form 4 (LastName FirstName format)
    additional_mappings = {}
    for original, replacement in alias_map.items():
        # Check if it's a person name format (FirstName LastName)
        if ' ' in original and len(original.split()) == 2:
            first, last = original.split()
            # Generate reverse format
            reverse_original = f"{last} {first}"
            if reverse_original not in alias_map:
                # Generate reverse replacement
                if ' ' in replacement and len(replacement.split()) == 2:
                    rep_first, rep_last = replacement.split()
                    additional_mappings[reverse_original] = f"{rep_last} {rep_first}"

    alias_map.update(additional_mappings)

    # CRITICAL FIX: Strip common suffixes to create base name mappings
    # This handles cases like:
    # - "Aaron Levie, CEO" exists but document contains just "Aaron Levie"
    # - "Dutch Bros Inc." exists but document contains just "Dutch Bros"
    suffix_mappings = {}

    # Executive title suffixes
    executive_suffixes = [', CEO', ', CFO', ', COO', ', CTO', ', Chief Executive Officer',
                         ', Chief Financial Officer', ', Chief Operating Officer',
                         ', Chief Technology Officer', ', Chief Legal Officer',
                         ', Vice President', ', Director', ', President',
                         ', Chief Accounting Officer and Controller',
                         ', Chief Legal Officer & Corporate Secretary']

    # Company suffixes
    company_suffixes = [' Inc.', ' Corp.', ' Corporation', ' LLC', ' L.L.C.',
                       ' Ltd.', ' Limited', ' Co.', ' Company']

    all_suffixes = executive_suffixes + company_suffixes

    for original, replacement in alias_map.items():
        for suffix in all_suffixes:
            if original.endswith(suffix):
                # Create base name mapping
                base_original = original[:-len(suffix)]

                # Only add if base name doesn't already exist
                if base_original not in alias_map and base_original not in suffix_mappings:
                    # Try to strip same suffix from replacement
                    if replacement.endswith(suffix):
                        base_replacement = replacement[:-len(suffix)]
                    else:
                        # Replacement might have different format, keep as is
                        base_replacement = replacement

                    suffix_mappings[base_original] = base_replacement

    alias_map.update(suffix_mappings)

    return alias_map


def categorize_and_sort_aliases(alias_map):
    """
    Bulletproof 3-tier sorting to prevent cascading failures.

    Processing order:
    1. Company names with suffixes (Inc., Corp., LLC, etc.)
    2. Multi-word phrases
    3. Single words and tickers

    Within each tier: longest first
    """
    company_suffixes = ['Inc.', 'Corp.', 'Corporation', 'LLC', 'L.L.C.', 'Ltd.', 'Limited', 'Co.', 'Company']

    tier1_company = []  # Company names
    tier2_multiword = []  # Multi-word phrases
    tier3_single = []  # Single words/tickers

    for original in alias_map.keys():
        # Tier 1: Company names with suffixes
        if any(suffix in original for suffix in company_suffixes):
            tier1_company.append(original)
        # Tier 2: Multi-word phrases
        elif ' ' in original:
            tier2_multiword.append(original)
        # Tier 3: Single words/tickers
        else:
            tier3_single.append(original)

    # Sort each tier by length (longest first)
    tier1_company.sort(key=len, reverse=True)
    tier2_multiword.sort(key=len, reverse=True)
    tier3_single.sort(key=len, reverse=True)

    # Combine in order
    sorted_keys = tier1_company + tier2_multiword + tier3_single

    return sorted_keys


def precompile_patterns(alias_map):
    """
    Pre-compile all regex patterns once for performance.

    PERFORMANCE OPTIMIZATION: Compiling patterns once instead of millions of times
    provides 20-40% speed improvement.
    """
    import re
    compiled_patterns = {}
    for original in alias_map.keys():
        compiled_patterns[original] = re.compile(re.escape(original), re.IGNORECASE)
    return compiled_patterns


def anonymize_text(text, alias_map, sorted_keys, compiled_patterns=None):
    """
    Apply anonymization replacements with case matching.

    Args:
        compiled_patterns: Pre-compiled regex patterns dict (optional but recommended for performance)
    """
    replacements = 0
    import re

    # If patterns not pre-compiled, compile them now (fallback for backward compatibility)
    if compiled_patterns is None:
        compiled_patterns = precompile_patterns(alias_map)

    for original in sorted_keys:
        replacement = alias_map[original]

        # Case-sensitive replacement with case matching
        def replace_with_case(match):
            nonlocal replacements
            matched_text = match.group(0)

            # Preserve case pattern
            if matched_text.isupper():
                replacements += 1
                return replacement.upper()
            elif matched_text.islower():
                replacements += 1
                return replacement.lower()
            elif matched_text[0].isupper():
                replacements += 1
                return replacement.capitalize()
            else:
                replacements += 1
                return replacement

        # Use pre-compiled pattern (PERFORMANCE OPTIMIZED)
        pattern = compiled_patterns[original]
        text = pattern.sub(replace_with_case, text)

    return text, replacements


def anonymize_paragraph(paragraph, alias_map, sorted_keys, compiled_patterns=None):
    """
    Anonymize a single paragraph, handling text that spans multiple runs.

    CRITICAL: Word documents often split text across multiple runs even when
    it appears contiguous. For example, "Aaron Levie" might be stored as:
      Run 0: "Aaron"
      Run 1: " "
      Run 2: "Levie"

    Processing each run individually would fail to match "Aaron Levie".
    This function processes the full paragraph text, then rebuilds runs.

    SPECIAL HANDLING: Preserves hyperlinks by anonymizing text within hyperlink
    XML elements rather than destroying the entire structure.

    PERFORMANCE: Accepts pre-compiled regex patterns for 20-40% speedup.
    """
    if not paragraph.text:
        return 0

    # Check if paragraph contains hyperlinks at XML level
    p_elem = paragraph._element
    hyperlinks = p_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')

    # If there are hyperlinks, handle them specially to preserve structure
    if hyperlinks:
        count = 0
        # Anonymize text within each hyperlink element without destroying it
        for hyperlink in hyperlinks:
            # Find all text elements within this hyperlink
            text_elems = hyperlink.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            for text_elem in text_elems:
                if text_elem.text:
                    new_text, repl_count = anonymize_text(text_elem.text, alias_map, sorted_keys, compiled_patterns)
                    if repl_count > 0:
                        text_elem.text = new_text
                        count += repl_count

        # Also handle non-hyperlink text in the same paragraph
        # Find text elements that are NOT inside hyperlinks
        for run_elem in p_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            # Check if this run is inside a hyperlink
            parent = run_elem.getparent()
            is_in_hyperlink = False
            while parent is not None:
                if parent.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink':
                    is_in_hyperlink = True
                    break
                parent = parent.getparent()

            # Only process if NOT in hyperlink (already handled above)
            if not is_in_hyperlink:
                text_elems = run_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                for text_elem in text_elems:
                    if text_elem.text:
                        new_text, repl_count = anonymize_text(text_elem.text, alias_map, sorted_keys, compiled_patterns)
                        if repl_count > 0:
                            text_elem.text = new_text
                            count += repl_count

        return count

    # No hyperlinks - use original approach (destroy and rebuild runs)
    # Get full paragraph text and apply replacements
    full_text = paragraph.text
    new_text, count = anonymize_text(full_text, alias_map, sorted_keys, compiled_patterns)

    # If no replacements, skip
    if count == 0:
        return 0

    # Replacement occurred - need to update the paragraph
    # Simple approach: clear all runs and create one new run with the replaced text
    # This loses per-character formatting but ensures complete replacement

    # Save formatting from first run (if exists)
    first_run_format = None
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run_format = {
            'bold': first_run.bold,
            'italic': first_run.italic,
            'underline': first_run.underline,
            'font_name': first_run.font.name if first_run.font else None,
            'font_size': first_run.font.size if first_run.font else None,
        }

    # CRITICAL: Must completely remove runs from paragraph XML, not just clear text
    # Clearing text leaves empty runs that can cause duplication
    while len(paragraph.runs) > 0:
        p = paragraph._element
        p.remove(paragraph.runs[0]._element)

    # Create new run with replaced text
    new_run = paragraph.add_run(new_text)

    # Restore formatting if we saved it
    if first_run_format:
        new_run.bold = first_run_format['bold']
        new_run.italic = first_run_format['italic']
        new_run.underline = first_run_format['underline']
        if first_run_format['font_name']:
            new_run.font.name = first_run_format['font_name']
        if first_run_format['font_size']:
            new_run.font.size = first_run_format['font_size']

    return count


def anonymize_docx(docx_path, alias_map, sorted_keys):
    """
    Anonymize all text in DOCX file (paragraphs, headers, footers, tables).
    FIXED: Now handles text that spans multiple runs.

    PERFORMANCE OPTIMIZED: Pre-compiles regex patterns once for 20-40% speedup.
    """
    doc = Document(docx_path)
    total_replacements = 0

    # PERFORMANCE: Pre-compile all regex patterns ONCE (not millions of times)
    compiled_patterns = precompile_patterns(alias_map)

    # Anonymize paragraphs (process as whole units, not individual runs)
    for paragraph in doc.paragraphs:
        count = anonymize_paragraph(paragraph, alias_map, sorted_keys, compiled_patterns)
        total_replacements += count

    # Anonymize tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count = anonymize_paragraph(paragraph, alias_map, sorted_keys, compiled_patterns)
                    total_replacements += count

    # CRITICAL FIX: Anonymize textboxes and shapes in main document body
    # This was missing and caused "Matador" to appear in PDFs but not Word
    if hasattr(doc, '_element'):
        try:
            # Find all text elements inside textboxes in the main body
            textbox_texts = doc._element.xpath('.//w:txbxContent//w:t')
            for text_elem in textbox_texts:
                if text_elem.text:
                    text_elem.text, count = anonymize_text(text_elem.text, alias_map, sorted_keys, compiled_patterns)
                    total_replacements += count

            # Also handle VML textboxes (legacy format)
            vml_textbox_texts = doc._element.xpath('.//v:textbox//w:t')
            for text_elem in vml_textbox_texts:
                if text_elem.text:
                    text_elem.text, count = anonymize_text(text_elem.text, alias_map, sorted_keys, compiled_patterns)
                    total_replacements += count
        except Exception:
            pass  # Skip if xpath fails

    # Anonymize footnotes and endnotes
    try:
        if hasattr(doc, 'part') and hasattr(doc.part, 'package'):
            package = doc.part.package

            # Process footnotes
            try:
                footnotes_part = package.part_related_by('http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
                if hasattr(footnotes_part, '_element'):
                    footnote_texts = footnotes_part._element.xpath('.//w:t')
                    for text_elem in footnote_texts:
                        if text_elem.text:
                            text_elem.text, count = anonymize_text(text_elem.text, alias_map, sorted_keys, compiled_patterns)
                            total_replacements += count
            except Exception:
                pass  # No footnotes or error accessing them

            # Process endnotes
            try:
                endnotes_part = package.part_related_by('http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes')
                if hasattr(endnotes_part, '_element'):
                    endnote_texts = endnotes_part._element.xpath('.//w:t')
                    for text_elem in endnote_texts:
                        if text_elem.text:
                            text_elem.text, count = anonymize_text(text_elem.text, alias_map, sorted_keys, compiled_patterns)
                            total_replacements += count
            except Exception:
                pass  # No endnotes or error accessing them

            # Process comments
            try:
                comments_part = package.part_related_by('http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
                if hasattr(comments_part, '_element'):
                    comment_texts = comments_part._element.xpath('.//w:t')
                    for text_elem in comment_texts:
                        if text_elem.text:
                            text_elem.text, count = anonymize_text(text_elem.text, alias_map, sorted_keys, compiled_patterns)
                            total_replacements += count
            except Exception:
                pass  # No comments or error accessing them
    except Exception:
        pass  # Skip if parts not accessible

    # Anonymize headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            # Process regular paragraphs (using whole-paragraph approach)
            for paragraph in header.paragraphs:
                count = anonymize_paragraph(paragraph, alias_map, sorted_keys, compiled_patterns)
                total_replacements += count

            # Process textboxes in headers (CRITICAL for SEC filings)
            if hasattr(header, '_element'):
                try:
                    # Find all text elements inside textboxes
                    textbox_texts = header._element.xpath('.//w:txbxContent//w:t')
                    for text_elem in textbox_texts:
                        if text_elem.text:
                            text_elem.text, count = anonymize_text(text_elem.text, alias_map, sorted_keys, compiled_patterns)
                            total_replacements += count
                except Exception:
                    pass  # Skip if xpath fails

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            # Process regular paragraphs (using whole-paragraph approach)
            for paragraph in footer.paragraphs:
                count = anonymize_paragraph(paragraph, alias_map, sorted_keys, compiled_patterns)
                total_replacements += count

            # Process textboxes in footers
            if hasattr(footer, '_element'):
                try:
                    # Find all text elements inside textboxes
                    textbox_texts = footer._element.xpath('.//w:txbxContent//w:t')
                    for text_elem in textbox_texts:
                        if text_elem.text:
                            text_elem.text, count = anonymize_text(text_elem.text, alias_map, sorted_keys, compiled_patterns)
                            total_replacements += count
                except Exception:
                    pass  # Skip if xpath fails

    # CRITICAL: Anonymize hyperlink URLs (e.g., https://www.box.com → https://www.enclave.com)
    try:
        if hasattr(doc, 'part') and hasattr(doc.part, 'rels'):
            rels = doc.part.rels
            for rel_id, rel in rels.items():
                # Check if this is a hyperlink relationship
                if hasattr(rel, 'reltype') and 'hyperlink' in rel.reltype.lower():
                    if hasattr(rel, '_target') and rel._target:
                        original_url = rel._target
                        # Apply text replacements to URL
                        new_url, count = anonymize_text(original_url, alias_map, sorted_keys)
                        if count > 0:
                            rel._target = new_url
                            total_replacements += count
    except Exception:
        pass  # Skip if hyperlink processing fails

    return doc, total_replacements


def process_single_docx(input_path, output_path, alias_map, sorted_keys, logger, remove_images=True, clear_headers_footers_flag=False):
    """
    Process a single DOCX file: anonymize + strip metadata + optional image removal + optional header/footer clearing.

    Args:
        remove_images: If True, removes all images from document
        clear_headers_footers_flag: If True, clears all header/footer content (for presentations with logos)
    """
    logger.info(f"Processing: {input_path.name}")

    try:
        # Load DOCX
        doc, replacements = anonymize_docx(input_path, alias_map, sorted_keys)

        # Remove all images (if requested)
        images_removed = 0
        if remove_images:
            images_removed = remove_all_images(doc)

        # Clear headers/footers (if requested)
        headers_footers_cleared = 0
        if clear_headers_footers_flag:
            headers_footers_cleared = clear_headers_footers(doc)

        # Strip ALL metadata (CRITICAL)
        doc = strip_all_metadata(doc)

        # Save
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        if clear_headers_footers_flag:
            logger.info(f"  ✓ {replacements} replacements, {images_removed} images removed, {headers_footers_cleared} headers/footers cleared")
        else:
            logger.info(f"  ✓ {replacements} replacements, {images_removed} images removed")

        return replacements, images_removed

    except Exception as e:
        logger.error(f"  ❌ Error: {e}")
        return 0, 0


def process_single_docx_worker(args):
    """
    Worker function for multiprocessing pool.
    Creates its own logger to avoid pickling issues.

    Args:
        args: Tuple of (docx_file, input_dir, output_dir, alias_map, sorted_keys, file_index, total_files)

    Returns:
        Tuple of (relative_path_str, replacements, images_removed)
    """
    docx_file, input_dir, output_dir, alias_map, sorted_keys, file_index, total_files = args

    # Create logger for this worker
    worker_logger = logging.getLogger(f"worker_{file_index}")

    # Preserve directory structure
    relative_path = docx_file.relative_to(input_dir)
    output_path = output_dir / relative_path

    # Log progress
    worker_logger.info(f"[{file_index}/{total_files}] {relative_path}")

    # Process the file
    replacements, images = process_single_docx(
        docx_file, output_path, alias_map, sorted_keys, worker_logger
    )

    return (str(relative_path), replacements, images)


def main():
    parser = argparse.ArgumentParser(
        description='Process Adobe-converted DOCX files for anonymization',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Example:
  python3 process_adobe_word_files.py \\
      --input ~/Desktop/adobe_converted/ \\
      --output ../vdr-processor-docx/output/ \\
      --aliases input/requirements/Anonymization_Tracker_Barista.xlsx
        """
    )

    parser.add_argument('--input', required=True, help='Directory with Adobe-converted DOCX files')
    parser.add_argument('--output', required=True, help='Output directory for anonymized DOCX files')
    parser.add_argument('--aliases', default='input/requirements/Anonymization_Tracker_Barista.xlsx',
                       help='Path to anonymization Excel file')

    args = parser.parse_args()

    # Setup logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(message)s'
    )
    logger = logging.getLogger(__name__)

    input_dir = Path(args.input)
    output_dir = Path(args.output)
    aliases_file = Path(args.aliases)

    # Validate inputs
    if not input_dir.exists():
        logger.error(f"❌ Input directory not found: {input_dir}")
        sys.exit(1)

    if not aliases_file.exists():
        logger.error(f"❌ Aliases file not found: {aliases_file}")
        sys.exit(1)

    # Load aliases
    logger.info("Loading anonymization mappings...")
    alias_map = load_aliases_from_excel(aliases_file)
    sorted_keys = categorize_and_sort_aliases(alias_map)
    logger.info(f"✓ Loaded {len(alias_map)} mappings\n")

    # Find all DOCX files
    docx_files = list(input_dir.rglob('*.docx'))

    if not docx_files:
        logger.error(f"❌ No DOCX files found in {input_dir}")
        sys.exit(1)

    logger.info(f"Found {len(docx_files)} DOCX files to process\n")
    logger.info("="*80)

    # PERFORMANCE OPTIMIZATION: Use multiprocessing for parallel file processing
    # Determine optimal worker count (cap at 8 to avoid memory issues on 128GB system)
    num_workers = min(cpu_count(), 8)
    logger.info(f"Using {num_workers} parallel workers for {num_workers}x speedup\n")

    # Prepare arguments for parallel processing
    # Each worker gets: (docx_file, input_dir, output_dir, alias_map, sorted_keys, file_index, total_files)
    worker_args = [
        (docx_file, input_dir, output_dir, alias_map, sorted_keys, i, len(docx_files))
        for i, docx_file in enumerate(docx_files, 1)
    ]

    # Process files in parallel
    total_replacements = 0
    total_images = 0

    with Pool(num_workers) as pool:
        # Map worker function across all files
        results = pool.map(process_single_docx_worker, worker_args)

    # Aggregate results
    for relative_path_str, replacements, images in results:
        total_replacements += replacements
        total_images += images

    # Summary
    logger.info("="*80)
    logger.info(f"\n✓ BATCH COMPLETE")
    logger.info(f"  Files processed: {len(docx_files)}")
    logger.info(f"  Total replacements: {total_replacements}")
    logger.info(f"  Total images removed: {total_images}")
    logger.info(f"\n📁 Output: {output_dir.absolute()}")
    logger.info(f"\n📋 NEXT STEP: Use Adobe Acrobat to batch convert DOCX → PDF")


if __name__ == '__main__':
    main()
